from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Ask how many flashcards
num = int(input("How many flashcards do you want to create? "))

# Collect flashcards
flashcards = []
for i in range(num):
    print(f"\nFlashcard {i+1}")
    question = input("Enter the question: ")
    answer = input("Enter the answer: ")
    flashcards.append({"question": question, "answer": answer})

# Create Word document
doc = Document()
doc.add_heading('Flashcards', 0)

for i, card in enumerate(flashcards, 1):
    # Question
    q_para = doc.add_paragraph()
    q_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    q_run = q_para.add_run(f"{i}. Q: {card['question']}")
    q_run.font.size = Pt(12)
    q_run.bold = True

    # Answer
    a_para = doc.add_paragraph()
    a_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    a_run = a_para.add_run(f"A: {card['answer']}")
    a_run.font.size = Pt(12)

    doc.add_paragraph("")  # Space between flashcards

# Save the document
filename = "flashcards.docx"
doc.save(filename)
print(f"\n✅ Flashcards saved to '{filename}'")
